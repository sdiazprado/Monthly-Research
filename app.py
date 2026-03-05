import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import calendar
import time

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual - Banxico", layout="wide")

# Inyección de CSS para cambiar el color al Azul Banxico (#00205B)
st.markdown("""
    <style>
    /* Estilo para los botones principales y de descarga */
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    /* Estilo para las etiquetas (tags) de los selectores múltiples (meses y años) */
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# FUNCIONES AUXILIARES (BACKEND)
# ==========================================

@st.cache_data(show_spinner="Descargando y procesando datos del BIS...")
def load_data_bis():
    url = "https://www.bis.org/api/document_lists/cbspeeches.json"
    response = requests.get(url)
    response.raise_for_status()
    data = response.json()

    speeches_dict = data.get("list", {})
    rows = []

    for path, speech in speeches_dict.items():
        title = html.unescape(speech.get("short_title", ""))
        date_str = speech.get("publication_start_date", "")

        link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")

        rows.append({
            "Date": date_str,
            "Title": title,
            "Link": link
        })

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    
    return df
@st.cache_data(show_spinner="Navegando y extrayendo discursos del BBk (Alemania)...")
def load_data_bbk(start_date_str, end_date_str):
    # El BBk espera las fechas en formato DD.MM.YYYY
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    rows = []
    page = 0
    
    while True:
        params = {
            'sort': 'bbksortdate desc',
            'dateFrom': start_date_str,
            'dateTo': end_date_str,
            'pageNumString': str(page)
        }
        
        try:
            response = requests.get(base_url, headers=headers, params=params)
            response.raise_for_status()
        except:
            break # Si la página bloquea o falla, salimos del bucle
            
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        
        if not items:
            break # Ya no hay resultados
            
        for item in items:
            # 1. Extraer Fecha
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            
            # 2. Extraer Enlace y Título
            data_div = item.find('div', class_='teasable__data')
            link = ""
            titulo = ""
            
            if data_div:
                a_tag = data_div.find('a', class_='teasable__link')
                if a_tag:
                    link = a_tag.get('href', '')
                    if link.startswith('/'):
                        link = "https://www.bundesbank.de" + link
                    
                    span_tag = a_tag.find('span', class_='link__label')
                    if span_tag:
                        titulo = span_tag.text.strip()
            
            if fecha_str and titulo:
                rows.append({
                    "Date": fecha_str,
                    "Title": titulo,
                    "Link": link
                })
                
        # Si la página trajo menos de 10 resultados, es la última página
        if len(items) < 10:
            break
            
        page += 1
        time.sleep(0.5) # Pausa amigable para no saturar al servidor alemán
        
    df = pd.DataFrame(rows)
    if not df.empty:
        # Convertir a datetime asumiendo el formato alemán DD.MM.YYYY
        df["Date"] = pd.to_datetime(df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
        
    return df
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Color Azul
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)

    # Subrayado
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    # Formato: Tamaño 12 (24 medios puntos)
    sz = docx.oxml.shared.OxmlElement('w:sz')
    sz.set(docx.oxml.shared.qn('w:val'), '24')
    rPr.append(sz)
    
    szCs = docx.oxml.shared.OxmlElement('w:szCs')
    szCs.set(docx.oxml.shared.qn('w:val'), '24')
    rPr.append(szCs)

    # Formato: Fuente Calibri
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri')
    rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_word(dataframe, title="Discursos", subtitle=""):
    doc = Document()
    
    # Título principal centrado
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo centrado (Rango de fechas)
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = 'Calibri'
        run_sub.font.size = Pt(12)

    # Añadir un espacio antes de la tabla
    doc.add_paragraph()

    # Tabla sin bordes
    table = doc.add_table(rows=1, cols=2)
    
    # Encabezados con formato Calibri 12
    hdr_cells = table.rows[0].cells
    for idx, header_text in enumerate(['Date', 'Title']):
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        # Dejamos los encabezados en negrita para diferenciar del contenido
        run.bold = True 

    # Llenado de datos (Sin negritas)
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        
        date_str = str(row['Date'])[:10]
        
        # Celda 1: Fecha normal (Calibri 12)
        p_date = row_cells[0].paragraphs[0]
        run_date = p_date.add_run(date_str)
        run_date.font.name = 'Calibri'
        run_date.font.size = Pt(12)
        
        # Celda 2: Título (El formato normal se inyecta por dentro de add_hyperlink)
        p_title = row_cells[1].paragraphs[0]
        add_hyperlink(p_title, str(row['Title']), str(row['Link']))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
# ==========================================
# INTERFAZ DE USUARIO (SIDEBAR Y NAVEGACIÓN)
# ==========================================

# 1. Logo institucional en la barra lateral
try:
    st.sidebar.image("logo_banxico.png", use_column_width=True)
except:
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")

# 2. Selector de Tipo de Documento
tipo_doc = st.sidebar.selectbox(
    "Selecciona el Tipo de Documento",
    ["Reportes", "Publicaciones Institucionales", "Investigación", "Discursos"]
)

# 3. Selector de Organismo (Listas exactas basadas en tus imágenes)
if tipo_doc == "Discursos":
    organismos = [
        "BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", 
        "BoC (Canadá)", "BoE (Inglaterra)", "BoJ (Japón)", "BPI", 
        "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"
    ]
elif tipo_doc == "Reportes":
    organismos = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
elif tipo_doc == "Investigación":
    organismos = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]
elif tipo_doc == "Publicaciones Institucionales":
    organismos = ["BM", "BPI", "CEF", "CEMLA", "FMI", "G20", "OCDE", "OEI"]

organismo_seleccionado = st.sidebar.selectbox("Selecciona el Organismo", organismos)

st.sidebar.markdown("---")
st.sidebar.info("Herramienta de extracción automatizada para la elaboración del boletín mensual.")


# ==========================================
# CONTENIDO PRINCIPAL (MAIN PAGE)
# ==========================================

st.title("Boletín Mensual de Organismos Internacionales")
st.markdown(f"**Explorador de {tipo_doc} - {organismo_seleccionado}**")
st.markdown("---")

# ==========================================
# MÓDULOS DE EXTRACCIÓN
# ==========================================

# MÓDULO: DISCURSOS -> BPI (Antes BIS)
if tipo_doc == "Discursos" and organismo_seleccionado == "BPI":
    
    st.subheader("1. Selecciona el Mes y Año")
    df = load_data_bis()

    anios_disponibles = df["Date"].dt.year.dropna().unique().tolist()
    anios_disponibles.sort(reverse=True)
    anios_str = [str(int(a)) for a in anios_disponibles]

    if "2026" in anios_str:
        anios_str.remove("2026")
        anios_str.insert(0, "2026")

    meses_dict = {
        "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4,
        "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8,
        "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
    }

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect(
            "Mes(es)",
            options=list(meses_dict.keys()),
            default=[] 
        )
    with col2:
        anios_seleccionados = st.multiselect(
            "Año(s)",
            options=anios_str,
            default=["2026"] if "2026" in anios_str else []
        )

    buscar = st.button("🔍 Buscar", type="primary")

    if buscar or "bis_df_filtrado" in st.session_state:
        
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año para realizar la búsqueda.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            mask = (df["Date"].dt.year.isin(anios_num)) & (df["Date"].dt.month.isin(meses_num))
            filtered_df = df[mask]
            
            st.session_state["bis_df_filtrado"] = filtered_df

            if len(filtered_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                
                # Columnas para el mensaje de éxito y el botón de descarga
                col_mensaje, col_boton = st.columns([3, 1])
                
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(filtered_df)}** discursos en **{str_meses} {str_anios}**.")
                
                with col_boton:
                    # Construimos el texto del subtítulo (ej. "Marzo 2026")
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    
                    # Llamamos a la función pasándole el nuevo parámetro 'subtitle'
                    word_file = generate_word(
                        filtered_df, 
                        title="BIS Central Bank Speeches", 
                        subtitle=subtitulo_fechas
                    )
                    st.download_button(
                        label="📄 Descargar en Word",
                        data=word_file,
                        file_name=f"bis_speeches_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )           
                
                # Procesar y mostrar la tabla en pantalla
                filtered_df_display = filtered_df.copy()
                filtered_df_display["Date"] = filtered_df_display["Date"].dt.strftime('%Y-%m-%d')
                filtered_df_display["Title"] = filtered_df_display.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1
                )

                st.markdown(
                    filtered_df_display[["Date", "Title"]].to_markdown(index=False),
                    unsafe_allow_html=True
                )

            else:
                st.warning("No hay discursos del BIS para las fechas seleccionadas. Intenta con otro mes.")
    else:
        st.info("👆 Selecciona el mes y año arriba y presiona **'Buscar'**.")
# MÓDULO: DISCURSOS -> BBk (Alemania)
elif tipo_doc == "Discursos" and organismo_seleccionado == "BBk (Alemania)":
    
    st.subheader("1. Selecciona el Mes y Año")

    # A diferencia del BIS, definimos los años fijos para no hacer scraping de todo el histórico
    anios_str = ["2026", "2025", "2024", "2023", "2022", "2021", "2020"]

    meses_dict = {
        "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4,
        "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8,
        "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
    }

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect(
            "Mes(es)",
            options=list(meses_dict.keys()),
            default=[] 
        )
    with col2:
        anios_seleccionados = st.multiselect(
            "Año(s)",
            options=anios_str,
            default=["2026"]
        )

    buscar = st.button("🔍 Buscar Discursos del BBk", type="primary")

    if buscar or "bbk_df_filtrado" in st.session_state:
        
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año para realizar la búsqueda.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            # Calculamos los extremos de fechas para pedirle solo lo necesario al servidor
            min_month, max_month = min(meses_num), max(meses_num)
            min_year, max_year = min(anios_num), max(anios_num)
            
            start_date_str = f"01.{min_month:02d}.{min_year}"
            last_day = calendar.monthrange(max_year, max_month)[1]
            end_date_str = f"{last_day:02d}.{max_month:02d}.{max_year}"

            # Llamamos a nuestra nueva función scraper
            df = load_data_bbk(start_date_str, end_date_str)
            
            if not df.empty:
                # Filtramos localmente para asegurar precisión (ej. si eligieron Enero y Marzo, descartamos Febrero)
                mask = (df["Date"].dt.year.isin(anios_num)) & (df["Date"].dt.month.isin(meses_num))
                filtered_df = df[mask]
            else:
                filtered_df = pd.DataFrame()
            
            st.session_state["bbk_df_filtrado"] = filtered_df

            if len(filtered_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                
                col_mensaje, col_boton = st.columns([3, 1])
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(filtered_df)}** discursos en **{str_meses} {str_anios}**.")
                
                with col_boton:
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    # Reutilizamos la función del Word que ya construimos antes
                    word_file = generate_word(
                        filtered_df, 
                        title="BBk Central Bank Speeches", 
                        subtitle=subtitulo_fechas
                    )
                    st.download_button(
                        label="📄 Descargar en Word",
                        data=word_file,
                        file_name=f"bbk_speeches_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                filtered_df_display = filtered_df.copy()
                filtered_df_display["Date"] = filtered_df_display["Date"].dt.strftime('%Y-%m-%d')
                filtered_df_display["Title"] = filtered_df_display.apply(
                    lambda x: f"[{x['Title']}]({x['Link']})", axis=1
                )

                st.markdown(
                    filtered_df_display[["Date", "Title"]].to_markdown(index=False),
                    unsafe_allow_html=True
                )
            else:
                st.warning("No hay discursos del BBk para las fechas seleccionadas.")

# MÓDULOS EN CONSTRUCCIÓN (Placeholder para el resto del menú)
else:
    st.info(f"El extractor de **{tipo_doc}** para **{organismo_seleccionado}** está en construcción.")
    st.write("Próximamente podrás extraer estos documentos de forma automatizada.")





